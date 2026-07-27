"""Сохранение сметы в Word (.docx) с колонтитулом."""

import os

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from models import Calculation


def save_calculation_to_docx(calc: Calculation, filepath: str) -> str:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.style.font.size = Pt(9)
    fp.add_run(f"Расчет выполнил: {calc.calculated_by}").italic = True

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("РАСЧЕТ СТОИМОСТИ ВЫЕЗДНЫХ РАБОТ")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x1A, 0x47, 0x8A)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(f"({calc.work_type})" if calc.work_type else "(поверка, калибровка, испытания)").font.size = Pt(12)

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    info.add_run(f"Дата: {calc.formatted_date()}\n")
    info.add_run(f"Сотрудник: {calc.role}\n")
    if calc.customer:
        info.add_run(f"Заказчик: {calc.customer}\n")
    if calc.work_location:
        info.add_run(f"Место проведения: {calc.work_location}\n")
    if calc.contract_number:
        info.add_run(f"Номер КП (договора): {calc.contract_number}\n")
    if calc.instrument_name:
        info.add_run(f"СИ: {calc.instrument_name}")

    doc.add_paragraph()

    table = doc.add_table(rows=0, cols=4)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = table.add_row().cells
    hdr_cells[0].text = "Наименование"
    hdr_cells[1].text = "Параметр"
    hdr_cells[2].text = "Стоимость"
    hdr_cells[3].text = "Сумма"
    for cell in hdr_cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True

    def add_row(name, param, price, total):
        row = table.add_row()
        row.cells[0].text = name
        row.cells[1].text = param
        row.cells[2].text = str(price)
        row.cells[3].text = str(total)

    add_row(
        "Командировочные",
        f"{calc.total_days} дн.",
        "",
        f"{calc.daily_cost} ₽",
    )

    add_row(
        "Суточные",
        f"{calc.total_days} дн.",
        "",
        f"{calc.daily_allowance_total} ₽",
    )

    add_row("Проезд", "", "", f"{calc.travel_cost} ₽")

    add_row("Проведение работ", "", "", f"{calc.work_cost} ₽")

    if calc.hotel_nights > 0:
        add_row(
            "Проживание",
            f"{calc.hotel_nights} ноч.",
            "",
            f"{calc.hotel_total} ₽",
        )

    info_row = table.add_row()
    info_row.cells[0].text = "Дней на работы"
    info_row.cells[1].text = str(calc.work_days)
    info_row.cells[2].text = "(информационно)"
    info_row.cells[3].text = "—"

    total_row = table.add_row()
    total_row.cells[0].text = "ИТОГО"
    total_row.cells[1].text = ""
    total_row.cells[2].text = ""
    total_row.cells[3].text = f"{calc.total} ₽"
    for cell in total_row.cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True

    doc.add_paragraph()

    os.makedirs(os.path.dirname(filepath) if os.path.dirname(filepath) else ".", exist_ok=True)
    doc.save(filepath)
    return filepath
